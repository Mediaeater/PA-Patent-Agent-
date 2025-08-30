# file: app/utils.py
from docx import Document
import xml.etree.ElementTree as ET

def basic_claim_checks(draft):
    """
    Performs basic quality checks on the draft claims.
    """
    issues = []
    claims = draft.get('claims', [])
    spec = (draft.get('detailed_description') or "").lower()
    for c in claims:
        cid = c['id']
        text = c['text'].lower()
        # 1) "comprising" check
        if "comprising" not in text and "consisting" not in text:
            issues.append({"claim_id":cid, "issue":"No 'comprising' or equivalent open language."})
        # 2) antecedent terms
        tokens = text.split()
        # collect simple nouns (rough heuristic)
        # check that each noun appears in spec (very rough)
        nouns = [w.strip(",.;()") for w in tokens if w.isalpha() and len(w)>3]
        missing = [n for n in set(nouns) if n not in spec]
        if missing:
            issues.append({"claim_id":cid, "issue":"Terms lacking support in spec", "terms": missing[:5]})
    return issues

def export_docx(final_json, path="patent_draft.docx"):
    """
    Exports the patent draft to a DOCX file.
    """
    doc = Document()
    meta = final_json.get('meta', {})
    doc.add_heading(final_json['draft']['title'] or meta.get('title','Patent Draft'), level=1)
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(final_json['draft']['abstract'])
    doc.add_heading('Background', level=2)
    doc.add_paragraph(final_json['draft']['background'])
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(final_json['draft']['summary'])
    doc.add_heading('Detailed Description', level=2)
    doc.add_paragraph(final_json['draft']['detailed_description'])
    doc.add_heading('Claims', level=2)
    for c in final_json['draft'].get('claims', []):
        doc.add_paragraph(f"Claim {c['id']}. {c['text']}")
    doc.save(path)
    return path

def export_simple_xml(final_json, path="patent.xml"):
    """
    Exports the patent draft to a simple XML file.
    """
    root = ET.Element("patent")
    meta = ET.SubElement(root, "meta")
    ET.SubElement(meta, "title").text = final_json['draft']['title']
    ET.SubElement(meta, "date").text = final_json['meta']['date']
    body = ET.SubElement(root, "body")
    ET.SubElement(body, "abstract").text = final_json['draft']['abstract']
    desc = ET.SubElement(body, "detailed_description")
    desc.text = final_json['draft']['detailed_description']
    claims = ET.SubElement(root, "claims")
    for c in final_json['draft']['claims']:
        claim_el = ET.SubElement(claims, "claim", id=str(c['id']))
        claim_el.text = c['text']
    tree = ET.ElementTree(root)
    tree.write(path, encoding="utf-8", xml_declaration=True)
    return path
